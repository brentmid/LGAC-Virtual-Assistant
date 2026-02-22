import logging
import re
from dataclasses import dataclass
from pathlib import Path

from .models import DocumentChunk

logger = logging.getLogger(__name__)


@dataclass
class IngestionMetrics:
    """Tracks quality metrics across the ingestion pipeline."""

    documents_processed: int = 0
    total_chunks: int = 0
    tables_detected: int = 0
    layout_tables_filtered: int = 0
    headers_removed: int = 0
    sections_detected: int = 0
    docx_tables_extracted: int = 0
    pdfplumber_fallbacks: int = 0

    def summary(self) -> str:
        lines = [
            f"Documents processed:     {self.documents_processed}",
            f"Total chunks:            {self.total_chunks}",
            f"Tables detected:         {self.tables_detected}",
            f"Layout tables filtered:  {self.layout_tables_filtered}",
            f"DOCX tables extracted:   {self.docx_tables_extracted}",
            f"Headers/footers removed: {self.headers_removed}",
            f"Sections detected:       {self.sections_detected}",
            f"pdfplumber fallbacks:    {self.pdfplumber_fallbacks}",
        ]
        return "\n".join(lines)


def _remove_repeated_headers(
    text: str, threshold: int = 3
) -> tuple[str, int]:
    """Remove repeated header/footer lines and decorative dividers.

    Lines (stripped, >= 10 chars) appearing more than `threshold` times are
    removed except for the first occurrence.  Duplicate decorative dividers
    (lines made entirely of whitespace / dashes / equals / underscores /
    asterisks / tildes / dots) are also collapsed to a single instance.

    Returns (cleaned_text, count_of_removed_lines).
    """
    divider_re = re.compile(r"^[\s\-=_*~.]+$")
    lines = text.split("\n")

    # Count occurrences of each normalised line
    counts: dict[str, int] = {}
    for line in lines:
        stripped = line.strip()
        if stripped:
            counts[stripped] = counts.get(stripped, 0) + 1

    seen: dict[str, int] = {}
    kept: list[str] = []
    removed = 0

    for line in lines:
        stripped = line.strip()

        # Decorative divider handling — keep only the first instance
        if stripped and divider_re.match(stripped):
            if stripped not in seen:
                seen[stripped] = 1
                kept.append(line)
            else:
                removed += 1
            continue

        # Repeated header/footer handling
        if stripped and len(stripped) >= 10 and counts.get(stripped, 0) > threshold:
            if stripped not in seen:
                seen[stripped] = 1
                kept.append(line)
            else:
                removed += 1
            continue

        kept.append(line)

    return "\n".join(kept), removed


def _reverse_cell_text(cell: str) -> str:
    """Reverse character order within each line and reverse line order.

    Fixes text from rotated PDF cells where characters are reversed
    and lines read bottom-to-top (e.g., 'esuohkaetS\\ns\\'remlaP' → "Palmer's Steakhouse").
    """
    lines = cell.strip().split("\n")
    reversed_lines = [line[::-1] for line in lines]
    reversed_lines.reverse()
    return " ".join(part.strip() for part in reversed_lines if part.strip())


def _has_reversed_text(cells: list[str]) -> bool:
    """Detect reversed text in a list of cell values.

    Checks if non-empty cells contain words (3+ alphabetic chars) that
    end with an uppercase letter and start with lowercase — the signature
    of a reversed proper noun (e.g., 'esuohkaetS' from 'Steakhouse').

    Returns True if a majority of qualifying words match the pattern.
    """
    if not cells:
        return False

    reversed_pattern = re.compile(r"^[a-z].*[A-Z]$")
    qualifying = 0
    matching = 0

    for cell in cells:
        for word in cell.replace("\n", " ").split():
            alpha_chars = sum(1 for c in word if c.isalpha())
            if alpha_chars >= 3:
                qualifying += 1
                if reversed_pattern.match(word):
                    matching += 1

    return qualifying > 0 and matching / qualifying > 0.5


def _extract_item_names(text: str) -> list[str]:
    """Extract item names from pdfplumber text containing Y/N grid data.

    Parses lines like 'Dress Hats (Women) Y Y Y Y Y Y Y Y Y Y'
    to extract 'Dress Hats (Women)'.
    """
    yn_values = {"Y", "N", "L*", "Y**", "Y***"}
    items = []
    for line in text.split("\n"):
        parts = line.strip().split()
        if len(parts) < 3:
            continue
        yn_count = 0
        for p in reversed(parts):
            if p in yn_values:
                yn_count += 1
            else:
                break
        if yn_count >= 5:
            name = " ".join(parts[: len(parts) - yn_count])
            if name:
                items.append(name)
    return items


def _format_reversed_table(
    table: list[list], text: str | None = None
) -> str:
    """Format a table with reversed headers into structured text.

    Fixes reversed column headers, filters empty columns, fills in missing
    row labels from extract_text() output, and produces structured output
    like 'Item | Venue1: Y | Venue2: N | ...'.
    """
    headers_row = table[0]

    # Find non-empty header columns (skip column 0 which is the row label)
    header_map: list[tuple[int, str]] = []
    for i in range(1, len(headers_row)):
        cell = headers_row[i]
        if cell and str(cell).strip():
            header_map.append((i, _reverse_cell_text(str(cell))))

    # Extract item names from text for filling missing row labels
    item_names = _extract_item_names(text) if text else []

    lines = []
    item_idx = 0
    for row in table[1:]:
        # Collect values from non-empty header columns
        vals = []
        for col_i, _ in header_map:
            v = str(row[col_i]).strip() if col_i < len(row) and row[col_i] else ""
            vals.append(v)

        if not any(vals):
            continue  # Skip empty rows

        # Get the label from the row or fall back to item_names
        label = str(row[0]).strip() if row[0] and str(row[0]).strip() else None
        if not label and item_idx < len(item_names):
            label = item_names[item_idx]
        if not label:
            label = "Unknown"

        item_idx += 1

        parts = [label]
        for (_, header_name), v in zip(header_map, vals):
            if v:
                parts.append(f"{header_name}: {v}")

        lines.append(" | ".join(parts))

    return "\n".join(lines)


def _has_orphaned_grid_data(text: str) -> bool:
    """Detect standalone Y/N lines indicating grid table data without column context.

    PyMuPDF sometimes extracts grid tables as individual Y/N values on
    separate lines, with column headers detached at the end of the text.
    """
    yn_values = {"Y", "N", "L*", "Y**"}
    yn_lines = 0
    for line in text.split("\n"):
        stripped = line.strip().rstrip("\t").strip()
        if stripped in yn_values:
            yn_lines += 1
    return yn_lines >= 20


def extract_text_pymupdf(path: Path) -> str:
    """Extract text from PDF using PyMuPDF."""
    import fitz

    doc = fitz.open(str(path))
    pages = []
    for page in doc:
        pages.append(page.get_text())
    doc.close()
    return "\n".join(pages)


def _is_layout_table(table: list[list]) -> bool:
    """Detect PDF layout-artifact tables that contain no meaningful data.

    Returns True when:
    - All header cells (first row) are empty/None
    - The table has only 1 column
    - More than 80% of all cells are empty
    """
    if not table or len(table) < 2:
        return True

    headers = table[0]

    # Single-column table
    if len(headers) <= 1:
        return True

    # All headers empty
    non_empty_headers = [c for c in headers if c and str(c).strip()]
    if not non_empty_headers:
        return True

    # >80% empty cells
    total_cells = 0
    empty_cells = 0
    for row in table:
        for cell in row:
            total_cells += 1
            if not cell or not str(cell).strip():
                empty_cells += 1

    if total_cells > 0 and empty_cells / total_cells > 0.8:
        return True

    return False


def _format_normal_table(table: list[list]) -> str:
    """Format a non-reversed PDF table with header context.

    Uses first row as headers and formats data rows as
    ``Header1: Value1 | Header2: Value2 | ...``.
    Skips columns with empty headers.  Returns empty string if fewer than
    2 non-empty header columns remain.
    """
    headers = table[0]

    # Build map of columns with non-empty headers
    header_map: list[tuple[int, str]] = []
    for i, h in enumerate(headers):
        if h and str(h).strip():
            header_map.append((i, str(h).strip()))

    if len(header_map) < 2:
        return ""

    lines: list[str] = []
    for row in table[1:]:
        parts: list[str] = []
        for col_i, header in header_map:
            val = str(row[col_i]).strip() if col_i < len(row) and row[col_i] else ""
            if val:
                parts.append(f"{header}: {val}")

        if parts:
            lines.append(" | ".join(parts))

    return "\n".join(lines)


def extract_text_pdfplumber(path: Path) -> str:
    """Extract text from PDF using pdfplumber (better for tables)."""
    import pdfplumber

    pages = []
    with pdfplumber.open(str(path)) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                pages.append(text)

            tables = page.extract_tables()
            for table in tables:
                if not table or len(table) < 2:
                    continue

                # Filter layout-artifact tables
                if _is_layout_table(table):
                    continue

                # Check first row for reversed text (rotated column headers)
                header_cells = [str(c) for c in table[0] if c and str(c).strip()]
                if header_cells and _has_reversed_text(header_cells):
                    pages.append(_format_reversed_table(table, text))
                else:
                    formatted = _format_normal_table(table)
                    if formatted:
                        pages.append(formatted)
    return "\n".join(pages)


def extract_pdf(path: Path) -> str:
    """Extract text from PDF, trying PyMuPDF first, falling back to pdfplumber."""
    text = extract_text_pymupdf(path)
    if len(text.strip()) < 100:
        logger.info(f"PyMuPDF yielded little text for {path.name}, trying pdfplumber")
        text = extract_text_pdfplumber(path)
    else:
        # Quality checks: detect garbled text that PyMuPDF can't handle well
        text_lines = [line.strip() for line in text.split("\n") if line.strip()]
        if _has_reversed_text(text_lines) or _has_orphaned_grid_data(text):
            logger.info(
                f"PyMuPDF output has quality issues for {path.name}, "
                "trying pdfplumber"
            )
            text = extract_text_pdfplumber(path)
    return text


def _format_docx_table(table) -> str:
    """Format a python-docx Table into structured text.

    Uses the first row as headers and formats subsequent rows as
    ``Header1: Value1 | Header2: Value2 | ...``.  Skips entirely-empty rows
    and deduplicates merged-cell text.
    """
    rows = table.rows
    if len(rows) < 2:
        return ""

    headers = [cell.text.strip() for cell in rows[0].cells]
    # Deduplicate adjacent merged-cell headers
    deduped: list[tuple[int, str]] = []
    prev = None
    for i, h in enumerate(headers):
        if h != prev:
            deduped.append((i, h))
            prev = h

    lines: list[str] = []
    for row in list(rows)[1:]:
        cells = [cell.text.strip() for cell in row.cells]
        # Deduplicate adjacent merged-cell values
        deduped_cells: list[tuple[int, str]] = []
        prev_val = None
        for i, v in enumerate(cells):
            if v != prev_val:
                deduped_cells.append((i, v))
                prev_val = v

        if not any(v for _, v in deduped_cells):
            continue  # skip empty rows

        parts: list[str] = []
        for col_i, header in deduped:
            # Find the value for this column index
            val = ""
            for ci, cv in deduped_cells:
                if ci == col_i:
                    val = cv
                    break
            if header and val:
                parts.append(f"{header}: {val}")
            elif val:
                parts.append(val)

        if parts:
            lines.append(" | ".join(parts))

    return "\n".join(lines)


def extract_docx(path: Path) -> str:
    """Extract text from DOCX file, including tables in document order."""
    import docx
    from docx.oxml.ns import qn

    doc = docx.Document(str(path))
    parts: list[str] = []
    tables_extracted = 0

    # Walk the document body children in order so paragraphs and tables
    # stay interleaved as they appear in the document.
    table_elements = {tbl._element: tbl for tbl in doc.tables}

    for child in doc.element.body:
        if child.tag == qn("w:p"):
            text = child.text
            # child.text only gets direct text; use the full paragraph API
            # by finding the matching Paragraph object
            for p in doc.paragraphs:
                if p._element is child:
                    text = p.text
                    break
            if text and text.strip():
                parts.append(text)
        elif child.tag == qn("w:tbl"):
            tbl_obj = table_elements.get(child)
            if tbl_obj:
                formatted = _format_docx_table(tbl_obj)
                if formatted:
                    parts.append(formatted)
                    tables_extracted += 1

    if tables_extracted:
        logger.info(f"  Extracted {tables_extracted} table(s) from DOCX")

    return "\n".join(parts)


def extract_document(path: Path) -> str:
    """Extract text from a document based on its extension."""
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return extract_pdf(path)
    elif suffix in (".docx", ".doc"):
        return extract_docx(path)
    else:
        raise ValueError(f"Unsupported file type: {suffix}")


def _find_section_headings(text: str) -> list[tuple[int, str]]:
    """Find section headings and their character positions in text.

    Detects:
    - ALL-CAPS lines (3+ alpha chars, <80 chars, not all digits/punctuation)
    - ``SECTION I/II/III/...`` patterns

    Returns list of ``(char_position, heading_text)`` sorted by position.
    """
    headings: list[tuple[int, str]] = []
    caps_re = re.compile(r"^[A-Z][A-Z\s\d\-/&,.'()]+$")
    section_re = re.compile(r"^SECTION\s+[IVXLCDM\d]+", re.IGNORECASE)

    pos = 0
    for line in text.split("\n"):
        stripped = line.strip()
        if stripped and len(stripped) < 80:
            alpha_count = sum(1 for c in stripped if c.isalpha())
            if alpha_count >= 3 and (
                caps_re.match(stripped) or section_re.match(stripped)
            ):
                headings.append((pos, stripped))
        pos += len(line) + 1  # +1 for the newline

    return headings


def chunk_text(
    text: str,
    source_name: str,
    chunk_size: int = 800,
    chunk_overlap: int = 100,
) -> list[DocumentChunk]:
    """Split text into overlapping chunks by approximate token count.

    Uses character-based chunking (~4 chars per token).  Tries to split at
    paragraph, sentence, or section-heading boundaries.  Prepends the most
    recent section heading to chunks that don't already contain it.
    """
    if not text.strip():
        return []

    # Approximate: ~4 chars per token
    char_limit = chunk_size * 4
    char_overlap = chunk_overlap * 4

    # Clean up text
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = text.strip()

    # Build section heading index
    headings = _find_section_headings(text)

    chunks = []
    start = 0

    while start < len(text):
        end = start + char_limit

        if end >= len(text):
            chunk_text_str = text[start:]
        else:
            # Try to break just before a heading line
            heading_break = -1
            for h_pos, _ in headings:
                if start + char_limit // 2 <= h_pos < end:
                    heading_break = h_pos
                    break  # use the first heading in the window

            if heading_break != -1:
                end = heading_break
            else:
                # Try to break at paragraph boundary
                para_break = text.rfind("\n\n", start + char_limit // 2, end)
                if para_break != -1:
                    end = para_break
                else:
                    # Try sentence boundary
                    sentence_break = text.rfind(
                        ". ", start + char_limit // 2, end
                    )
                    if sentence_break != -1:
                        end = sentence_break + 1

            chunk_text_str = text[start:end]

        chunk_text_str = chunk_text_str.strip()
        if chunk_text_str:
            # Find the most recent heading before this chunk's start
            current_heading = None
            for h_pos, h_text in headings:
                if h_pos <= start:
                    current_heading = h_text
                else:
                    break

            # Prepend heading if chunk doesn't already start with it
            if current_heading and not chunk_text_str.startswith(
                current_heading
            ):
                chunk_text_str = f"[{current_heading}]\n{chunk_text_str}"

            metadata: dict = {
                "source": source_name,
                "chunk_index": len(chunks),
            }
            if current_heading:
                metadata["section"] = current_heading

            chunks.append(
                DocumentChunk(
                    text=chunk_text_str,
                    metadata=metadata,
                )
            )

        if end >= len(text):
            break

        start = end - char_overlap

    return chunks


def ingest_documents(
    docs_dir: Path,
    chunk_size: int = 800,
    chunk_overlap: int = 100,
) -> tuple[list[DocumentChunk], IngestionMetrics]:
    """Process all documents in a directory into chunks.

    Returns a tuple of (chunks, metrics).
    """
    all_chunks: list[DocumentChunk] = []
    metrics = IngestionMetrics()
    supported = (".pdf", ".docx", ".doc")

    if not docs_dir.exists():
        logger.warning(f"Documents directory not found: {docs_dir}")
        return all_chunks, metrics

    files = sorted(f for f in docs_dir.iterdir() if f.suffix.lower() in supported)
    logger.info(f"Found {len(files)} documents to process")

    for filepath in files:
        try:
            logger.info(f"Processing: {filepath.name}")
            text = extract_document(filepath)

            # Remove repeated headers/footers
            text, removed = _remove_repeated_headers(text)
            if removed:
                logger.info(f"  Removed {removed} repeated header/footer lines")
                metrics.headers_removed += removed

            # Count sections
            headings = _find_section_headings(text)
            metrics.sections_detected += len(headings)

            chunks = chunk_text(text, filepath.name, chunk_size, chunk_overlap)
            logger.info(f"  -> {len(chunks)} chunks from {filepath.name}")
            all_chunks.extend(chunks)
            metrics.documents_processed += 1
        except Exception as e:
            logger.error(f"Failed to process {filepath.name}: {e}")

    metrics.total_chunks = len(all_chunks)
    logger.info(f"Total: {len(all_chunks)} chunks from {len(files)} documents")
    return all_chunks, metrics
